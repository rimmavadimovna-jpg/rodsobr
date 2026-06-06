"""Экспорт всего банка в редактируемый Word-файл для ручной выверки.

Запуск:  python -m letovo_bot.data.export_docx [--db путь] [--out файл.docx]

Для каждого задания выводит: условие (как видит ученик), правильный ответ/эталон,
источник и пустое поле «КОММЕНТАРИЙ / ИСПРАВЛЕНИЕ» для правок преподавателя.
Указывается ID задания в банке — по нему правки переносятся обратно в build_bank.py.
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path

from .. import config
from ..core import db
from ..core.models import Task, TaskType

TYPE_NAMES = {
    1: "«Третий лишний» по правописанию",
    2: "Глаголы заданного спряжения",
    3: "Схемы предложений",
    4: "Пунктуация + грамматические основы + объяснение",
    5: "Конструирование предложений",
    6: "Исправление грамматических ошибок",
    7: "Фонетика: подсчёт звука",
    8: "Синонимы с разными корнями",
    9: "Словообразовательная цепочка + морфемика",
    10: "«Четвёртое лишнее» по морфологии",
    11: "Анализ текста: верные/ошибочные утверждения",
    12: "Фразеологизм в текст + толкование",
}


# --------------------------------------------------------------------------- #
# Текстовое представление условия и эталона (по типам)
# --------------------------------------------------------------------------- #
def render_condition(task: Task) -> list[str]:
    p = task.payload
    tt = TaskType(task.task_type)
    out = [p.get("instruction", "")]
    if tt == TaskType.THIRD_EXTRA:
        for i, r in enumerate(p["rows"], 1):
            out.append(f"  Ряд {i}: {', '.join(r['words'])}   (принцип: {r.get('principle','')})")
    elif tt == TaskType.CONJUGATION:
        out.append(f"  Формы: {', '.join(p['forms'])}")
        out.append(f"  Целевое спряжение: {p.get('target_conjugation')}")
    elif tt in (TaskType.SCHEMES, TaskType.PUNCTUATION):
        for i, s in enumerate(p["sentences"], 1):
            out.append(f"  {i}) {s}")
    elif tt == TaskType.CONSTRUCT:
        out.append("  Слова: " + "; ".join(f"{w['word']} — {w['meaning']}" for w in p["words"]))
        out.append("  Фразеологизмы: " + "; ".join(f"{x['phraseme']} — {x['meaning']}" for x in p["phrasemes"]))
        out.append(f"  Требования: {p.get('requirements','')}")
    elif tt == TaskType.GRAMMAR_FIX:
        for i, s in enumerate(p["sentences"], 1):
            out.append(f"  {i}) {s}")
    elif tt == TaskType.PHONETICS:
        out.append(f"  Предложение: «{p['sentence']}»   Звук: {p['sound']}")
    elif tt == TaskType.SYNONYMS:
        out.append(f"  Слово: {p['word']}   Контекст: «{p.get('context','')}»")
    elif tt == TaskType.WORD_FORMATION:
        out.append(f"  Слова (вперемешку): {', '.join(p['words'])}")
        out.append(f"  Разобрать слово: {p['target_word']}")
    elif tt == TaskType.FOURTH_EXTRA:
        for i, r in enumerate(p["rows"], 1):
            out.append(f"  Ряд {i}: {', '.join(r['words'])}")
    elif tt == TaskType.TEXT_STATEMENTS:
        out.append(f"  Текст: {p['text']}")
        out.append(f"  Вопрос: указать {p.get('ask','')} утверждения")
        for i, s in enumerate(p["statements"], 1):
            out.append(f"  {i}) {s}")
    elif tt == TaskType.PHRASEME:
        out.append(f"  Текст: {p['text']}")
        out.append(f"  Абзац № {p.get('paragraph')}")
    return [x for x in out if x]


def render_answer(task: Task) -> list[str]:
    a = task.answer
    tt = TaskType(task.task_type)
    out: list[str] = []
    if tt == TaskType.THIRD_EXTRA:
        for i, r in enumerate(a["rows"], 1):
            props = r.get("props")
            extra = f"  Ряд {i}: лишнее «{r['extra']}» → пишется «{r.get('spelling','')}»"
            if props:
                extra += f"   (признаки слов: {props})"
            out.append(extra)
    elif tt == TaskType.CONJUGATION:
        out.append(f"  Выписать формы: {', '.join(a['expected_forms'])}")
        out.append(f"  Инфинитивы: {', '.join(a['expected_infinitives'])}")
        out.append(f"  Отвлекающие (не выписывать): {', '.join(a.get('distractor_forms', []))}")
    elif tt == TaskType.SCHEMES:
        for i, it in enumerate(a["items"], 1):
            eq = it.get("equivalents", [])
            out.append(f"  {i}) {it['canonical']}" + (f"   (равнозначно: {', '.join(eq)})" if eq else ""))
    elif tt == TaskType.PUNCTUATION:
        for i, it in enumerate(a["items"], 1):
            out.append(f"  {i}) {it['reference']}")
            out.append(f"       основы: {it.get('bases')}")
            out.append(f"       правило: {it.get('rule','')}")
    elif tt == TaskType.CONSTRUCT:
        out.append(f"  Образец: {a.get('example','')}")
    elif tt == TaskType.GRAMMAR_FIX:
        out.append(f"  Ошибочные предложения: {a['wrong']}")
        for num, fx in a.get("fixes", {}).items():
            out.append(f"  {num}) исправлено: {fx['corrected']}")
            if fx.get("correct_fragments"):
                out.append(f"       допустимо: {fx['correct_fragments']}; неверный фрагмент: «{fx.get('wrong_fragment','')}»")
    elif tt == TaskType.PHONETICS:
        out.append(f"  Количество: {a['count']}")
        out.append(f"  Объяснение: {a.get('explanation','')}")
        if a.get("words_with_sound"):
            out.append(f"  Слова со звуком: {', '.join(a['words_with_sound'])}")
    elif tt == TaskType.SYNONYMS:
        for x in a["allowed"]:
            out.append(f"  • {x['syn']}  (корень: {x.get('root','?')}, стиль: {x.get('style','')})")
    elif tt == TaskType.WORD_FORMATION:
        out.append(f"  Цепочка: {' → '.join(a['chain'])}")
        out.append(f"  Способ: {a.get('method','')}")
        m = a.get("morphemes", {})
        out.append(f"  Разбор: приставка={m.get('prefix','')!r} корень={m.get('root','')!r} "
                   f"суффикс={m.get('suffix','')!r} окончание={m.get('ending','')!r}")
        for st in a.get("steps", []):
            out.append(f"       {st['from']} → {st['to']}  (источник: {st.get('source','')})")
    elif tt == TaskType.FOURTH_EXTRA:
        for i, r in enumerate(a["rows"], 1):
            extra = f"  Ряд {i}: лишнее «{r['extra']}» — {r.get('feature','')}"
            if r.get("props"):
                extra += f"   (признаки: {r['props']})"
            out.append(extra)
    elif tt == TaskType.TEXT_STATEMENTS:
        out.append(f"  Ключ (номера): {a['key']}")
    elif tt == TaskType.PHRASEME:
        for x in a["allowed"]:
            out.append(f"  • {x['phraseme']} — {x['meaning']}   (подходит: {x.get('fits','')})")
    return out


# --------------------------------------------------------------------------- #
# Сборка .docx
# --------------------------------------------------------------------------- #
def export(db_path: str | Path, out_path: str | Path) -> None:
    import docx
    from docx.shared import Pt, RGBColor

    conn = db.connect(db_path)
    tasks = sorted(db.all_tasks(conn), key=lambda t: (int(t.task_type), t.id))

    doc = docx.Document()
    doc.add_heading("Банк заданий Letovo-бота — выверка", level=0)
    intro = doc.add_paragraph()
    intro.add_run(
        "Ниже — все задания из банка. Под каждым есть поле «КОММЕНТАРИЙ / ИСПРАВЛЕНИЕ» — "
        "впишите туда, что не так и как должно быть (можно прямо в этом файле, выделяя правки). "
        "ID задания указан в заголовке — по нему я перенесу ваши правки обратно в банк. "
        "Затем пришлите файл мне."
    ).italic = True

    current_type = None
    for t in tasks:
        if int(t.task_type) != current_type:
            current_type = int(t.task_type)
            doc.add_heading(f"Тип {current_type}. {TYPE_NAMES.get(current_type,'')}", level=1)

        h = doc.add_heading(level=2)
        h.add_run(f"Задание #{t.id}  (тема: {t.topic}, сложность: {t.difficulty})")

        doc.add_paragraph().add_run("УСЛОВИЕ (как видит ученик):").bold = True
        for line in render_condition(t):
            doc.add_paragraph(line)

        doc.add_paragraph().add_run("ПРАВИЛЬНЫЙ ОТВЕТ / ЭТАЛОН:").bold = True
        for line in render_answer(t):
            doc.add_paragraph(line)

        src = doc.add_paragraph()
        src.add_run("ИСТОЧНИК: ").bold = True
        src.add_run(t.source or "—")

        cm = doc.add_paragraph()
        run = cm.add_run("КОММЕНТАРИЙ / ИСПРАВЛЕНИЕ: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        cm.add_run("________________________________________________________")

        doc.add_paragraph("—" * 30)

    # тексты для зад. 11–12
    texts = conn.execute("SELECT id, body, license FROM texts").fetchall()
    if texts:
        doc.add_heading("Тексты для заданий 11–12 (собственные, не Летово)", level=1)
        for row in texts:
            ph = doc.add_paragraph()
            ph.add_run(f"Текст #{row['id']} (лицензия: {row['license']}): ").bold = True
            ph.add_run(row["body"])
            cm = doc.add_paragraph()
            r = cm.add_run("КОММЕНТАРИЙ / ИСПРАВЛЕНИЕ: ")
            r.bold = True
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            cm.add_run("________________________________________________________")

    doc.save(str(out_path))
    print(f"Готово: {out_path} (заданий: {len(tasks)})")
    conn.close()


def main() -> None:
    ap = argparse.ArgumentParser(description="Экспорт банка в Word для выверки")
    ap.add_argument("--db", default=str(config.BANK_PATH))
    ap.add_argument("--out", default="Банк_заданий_на_выверку.docx")
    args = ap.parse_args()
    export(args.db, args.out)


if __name__ == "__main__":
    main()
